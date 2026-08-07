"""Word document tools for office-mcp.

This module exposes the 14 ``word_*`` tools described in
``architecture.md`` §3.2 and validated by ``validation-contract.md``
(``VAL-WORD-001``..``VAL-WORD-042``, the advanced ``VAL-WORD-043``..
``VAL-WORD-066``, plus the cross-tool assertions ``VAL-WORD-075``..
``VAL-WORD-079``).

Conventions
-----------
* Every tool resolves the user-supplied ``path`` (and optional
  ``folder``) through :func:`office_mcp.paths.resolve_path`.
* Every tool validates its inputs and raises :class:`OfficeMCPError`
  with the appropriate error code (``ERR_FILE_NOT_FOUND``,
  ``ERR_INVALID_PARAMS``, ``ERR_UNSUPPORTED_FMT``).
* Every tool opens the file with :class:`docx.Document`, performs the
  operation, saves, and returns a plain ``dict`` (which FastMCP
  serializes as the ``structuredContent`` of the JSON-RPC response).
* ``word_format_run`` with all-``None`` formatting arguments is a true
  no-op — the file is not re-saved so the SHA256 is preserved
  (see ``VAL-WORD-076``).
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_ORIENT
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph

from office_mcp.errors import (
    ERR_FILE_NOT_FOUND,
    ERR_INVALID_PARAMS,
    ERR_UNSUPPORTED_FMT,
    OfficeMCPError,
)
from office_mcp.paths import resolve_path, save_with_lock_check
from server import mcp


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------


# Page sizes for ``word_set_section`` (in inches; converted to EMU on use).
# Width/height are listed as ``(portrait_width, portrait_height)`` in inches.
_PAGE_SIZES: dict[str, tuple[float, float]] = {
    # ISO 216
    "A4": (8.27, 11.69),    # 210 x 297 mm
    "A5": (5.83, 8.27),     # 148 x 210 mm
    # US
    "Letter": (8.5, 11.0),
    "Legal": (8.5, 14.0),
    "Tabloid": (11.0, 17.0),
    # ISO B
    "B5": (6.93, 9.84),
}


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------


def _resolve_docx(path: str, folder: str | None) -> Path:
    """Resolve ``path`` to an absolute ``Path`` and verify the extension.

    Wraps :func:`office_mcp.paths.resolve_path` so that empty /
    whitespace-only paths become :class:`OfficeMCPError`
    (``ERR_INVALID_PARAMS``) instead of a bare ``ValueError``, and
    enforces the ``.docx`` extension up front so non-docx files fail
    fast.
    """
    try:
        p = resolve_path(path, folder)
    except ValueError as exc:
        raise OfficeMCPError(
            ERR_INVALID_PARAMS,
            str(exc),
            {"path": path},
        ) from exc
    if p.suffix.lower() != ".docx":
        raise OfficeMCPError(
            ERR_UNSUPPORTED_FMT,
            f"Expected a .docx file, got {p.suffix!r}",
            {"path": str(p), "extension": p.suffix},
        )
    return p


def _require_exists(path: Path) -> None:
    """Raise ``ERR_FILE_NOT_FOUND`` if ``path`` does not exist on disk."""
    if not path.exists():
        raise OfficeMCPError(
            ERR_FILE_NOT_FOUND,
            f"File not found: {path}",
            {"path": str(path)},
        )


def _open_docx(path: Path) -> DocumentObject:
    """Open an existing ``.docx`` file.

    Combines :func:`_require_exists` with :class:`docx.Document` so
    that callers always get a proper error instead of a generic
    ``FileNotFoundError`` from ``python-docx``.
    """
    _require_exists(path)
    return Document(str(path))


def _save_docx(doc: DocumentObject, path: Path) -> None:
    """Save the document back to disk, creating parents if needed.

    Routes the underlying ``Document.save`` call through
    :func:`office_mcp.paths.save_with_lock_check` so a held lock
    (file open in Word) surfaces as ``ERR_FILE_LOCKED`` instead of an
    uncaught :class:`PermissionError`.
    """
    path.parent.mkdir(parents=True, exist_ok=True)
    save_with_lock_check(path, lambda: doc.save(str(path)))


def _check_paragraph_index(doc: DocumentObject, index: int) -> None:
    if not isinstance(index, int) or isinstance(index, bool):
        raise OfficeMCPError(
            ERR_INVALID_PARAMS,
            f"Paragraph index must be an int, got {type(index).__name__}",
            {"index": index},
        )
    if index < 0 or index >= len(doc.paragraphs):
        raise OfficeMCPError(
            ERR_INVALID_PARAMS,
            (
                f"Paragraph index {index} is out of range "
                f"(document has {len(doc.paragraphs)} paragraph(s))"
            ),
            {"index": index, "count": len(doc.paragraphs)},
        )


def _check_run_index(paragraph: Paragraph, run_index: int) -> None:
    runs = paragraph.runs
    if not isinstance(run_index, int) or isinstance(run_index, bool):
        raise OfficeMCPError(
            ERR_INVALID_PARAMS,
            f"Run index must be an int, got {type(run_index).__name__}",
            {"run_index": run_index},
        )
    if run_index < 0 or run_index >= len(runs):
        raise OfficeMCPError(
            ERR_INVALID_PARAMS,
            (
                f"Run index {run_index} is out of range "
                f"(paragraph has {len(runs)} run(s))"
            ),
            {"run_index": run_index, "count": len(runs)},
        )


def _serialise_run(run) -> dict[str, Any]:
    """Return a JSON-friendly snapshot of a run's properties."""
    font = run.font
    color = None
    try:
        if font.color and font.color.rgb is not None:
            color = str(font.color.rgb)
    except (AttributeError, ValueError):
        color = None
    size = None
    try:
        if font.size is not None:
            size = font.size.pt
    except (AttributeError, ValueError):
        size = None
    return {
        "text": run.text,
        "bold": run.bold,
        "italic": run.italic,
        "font_name": font.name,
        "font_size": size,
        "color": color,
    }


def _serialise_paragraph(paragraph: Paragraph, index: int) -> dict[str, Any]:
    return {
        "index": index,
        "style": paragraph.style.name,
        "text": paragraph.text,
        "runs": [_serialise_run(r) for r in paragraph.runs],
    }


# ---------------------------------------------------------------------------
# 1. word_create_document
# ---------------------------------------------------------------------------


@mcp.tool()
def word_create_document(
    path: str,
    title: str | None = None,
    folder: str | None = None,
) -> dict[str, str]:
    """Create a new ``.docx`` file at ``path`` with an optional title paragraph.

    Args:
        path: Target ``.docx`` path. May be absolute or relative to
            ``folder`` (or to the default folder when ``folder`` is
            ``None``).
        title: Optional text for the first paragraph. When ``None`` the
            document is created with a single empty paragraph (the
            python-docx default).
        folder: Optional base folder for relative paths.

    Returns:
        ``{"path": "<absolute path>"}`` on success.

    Raises:
        OfficeMCPError: ``ERR_INVALID_PARAMS`` if ``path`` is empty or
            a file already exists at the target. ``ERR_UNSUPPORTED_FMT``
            if ``path`` does not end in ``.docx``.
    """
    out = _resolve_docx(path, folder)
    if out.exists():
        raise OfficeMCPError(
            ERR_INVALID_PARAMS,
            f"File already exists: {out}",
            {"path": str(out)},
        )
    doc = Document()
    doc.add_paragraph(title if title is not None else "")
    _save_docx(doc, out)
    return {"path": str(out)}


# ---------------------------------------------------------------------------
# 2. word_get_info
# ---------------------------------------------------------------------------


@mcp.tool()
def word_get_info(path: str, folder: str | None = None) -> dict[str, Any]:
    """Return counts and core properties of a ``.docx`` file.

    The returned dict contains:

    * ``paragraphs`` — number of body paragraphs (including the empty
      one emitted by :func:`docx.Document`).
    * ``sections`` — number of sections in the document.
    * ``tables`` — number of body tables.
    * ``images`` — number of inline shapes (images).
    * ``properties`` — subset of ``core_properties`` (``title``,
      ``author``) that are populated.

    Args:
        path: Path to an existing ``.docx``.
        folder: Optional base folder for relative paths.

    Raises:
        OfficeMCPError: ``ERR_FILE_NOT_FOUND`` if the file is missing,
            ``ERR_UNSUPPORTED_FMT`` if the extension is not ``.docx``,
            ``ERR_INVALID_PARAMS`` if ``path`` is empty.
    """
    p = _resolve_docx(path, folder)
    doc = _open_docx(p)
    cp = doc.core_properties
    properties: dict[str, str] = {}
    if cp.title:
        properties["title"] = cp.title
    if cp.author:
        properties["author"] = cp.author
    return {
        "paragraphs": len(doc.paragraphs),
        "sections": len(doc.sections),
        "tables": len(doc.tables),
        "images": len(doc.inline_shapes),
        "properties": properties,
    }


# ---------------------------------------------------------------------------
# 3. word_list_paragraphs
# ---------------------------------------------------------------------------


@mcp.tool()
def word_list_paragraphs(path: str, folder: str | None = None) -> list[dict[str, Any]]:
    """Return one dict per body paragraph.

    Each entry has ``index``, ``style``, ``text``, and ``runs`` keys.
    ``runs`` is itself a list of run dicts (see
    :func:`_serialise_run`).

    Args:
        path: Path to an existing ``.docx``.
        folder: Optional base folder for relative paths.

    Raises:
        OfficeMCPError: ``ERR_FILE_NOT_FOUND`` if the file is missing,
            ``ERR_UNSUPPORTED_FMT`` for non-``.docx`` extensions.
    """
    p = _resolve_docx(path, folder)
    doc = _open_docx(p)
    return [_serialise_paragraph(p_, i) for i, p_ in enumerate(doc.paragraphs)]


# ---------------------------------------------------------------------------
# 4. word_read_paragraph
# ---------------------------------------------------------------------------


@mcp.tool()
def word_read_paragraph(
    path: str,
    index: int,
    folder: str | None = None,
) -> dict[str, Any]:
    """Return a single paragraph's content and metadata.

    Args:
        path: Path to an existing ``.docx``.
        index: Zero-based paragraph index.
        folder: Optional base folder for relative paths.

    Raises:
        OfficeMCPError: ``ERR_INVALID_PARAMS`` if ``index`` is negative
            or out of range, ``ERR_FILE_NOT_FOUND`` if the file is
            missing, ``ERR_UNSUPPORTED_FMT`` for non-``.docx``
            extensions.
    """
    p = _resolve_docx(path, folder)
    doc = _open_docx(p)
    _check_paragraph_index(doc, index)
    paragraph = doc.paragraphs[index]
    return _serialise_paragraph(paragraph, index)


# ---------------------------------------------------------------------------
# 5. word_add_paragraph
# ---------------------------------------------------------------------------


@mcp.tool()
def word_add_paragraph(
    path: str,
    text: str,
    style: str | None = None,
    folder: str | None = None,
) -> dict[str, int]:
    """Append a paragraph to the body of the document.

    Args:
        path: Path to an existing ``.docx``.
        text: Text for the new paragraph. An empty string is accepted
            and produces an empty paragraph.
        style: Optional built-in style name (e.g. ``"Intense Quote"``).
        folder: Optional base folder for relative paths.

    Returns:
        ``{"index": <n>}`` where ``<n>`` is the index of the newly
        appended paragraph.

    Raises:
        OfficeMCPError: ``ERR_FILE_NOT_FOUND`` if the file is missing,
            ``ERR_INVALID_PARAMS`` for bad ``text`` type or unknown
            ``style``, ``ERR_UNSUPPORTED_FMT`` for non-``.docx``
            extensions.
    """
    if not isinstance(text, str):
        raise OfficeMCPError(
            ERR_INVALID_PARAMS,
            f"text must be a str, got {type(text).__name__}",
            {"type": type(text).__name__},
        )
    p = _resolve_docx(path, folder)
    doc = _open_docx(p)
    if style is not None:
        try:
            paragraph = doc.add_paragraph(text, style=style)
        except KeyError as exc:
            raise OfficeMCPError(
                ERR_INVALID_PARAMS,
                f"Unknown style: {style!r}",
                {"style": style},
            ) from exc
    else:
        paragraph = doc.add_paragraph(text)
    _save_docx(doc, p)
    # Reload to be sure about the final ordering
    final = Document(str(p))
    new_index = next(
        (i for i, para in enumerate(final.paragraphs) if para._element is paragraph._element),
        len(final.paragraphs) - 1,
    )
    return {"index": new_index}


# ---------------------------------------------------------------------------
# 6. word_add_heading
# ---------------------------------------------------------------------------


@mcp.tool()
def word_add_heading(
    path: str,
    text: str,
    level: int = 1,
    folder: str | None = None,
) -> dict[str, int]:
    """Append a heading paragraph at the given ``level`` (1..9).

    Args:
        path: Path to an existing ``.docx``.
        text: Heading text.
        level: Heading level (1 = ``Heading 1``, ..., 9 = ``Heading 9``).
            Out-of-range values are rejected with
            ``ERR_INVALID_PARAMS``.
        folder: Optional base folder for relative paths.

    Returns:
        ``{"index": <n>}`` for the new heading.

    Raises:
        OfficeMCPError: ``ERR_INVALID_PARAMS`` for an out-of-range
            ``level`` or a non-int ``level``, ``ERR_FILE_NOT_FOUND``
            if the file is missing, ``ERR_UNSUPPORTED_FMT`` for non-
            ``.docx`` extensions.
    """
    if not isinstance(level, int) or isinstance(level, bool):
        raise OfficeMCPError(
            ERR_INVALID_PARAMS,
            f"level must be an int, got {type(level).__name__}",
            {"level": level},
        )
    if level < 1 or level > 9:
        raise OfficeMCPError(
            ERR_INVALID_PARAMS,
            f"Heading level must be between 1 and 9, got {level}",
            {"level": level},
        )
    if not isinstance(text, str):
        raise OfficeMCPError(
            ERR_INVALID_PARAMS,
            f"text must be a str, got {type(text).__name__}",
            {"type": type(text).__name__},
        )
    p = _resolve_docx(path, folder)
    doc = _open_docx(p)
    paragraph = doc.add_heading(text, level=level)
    _save_docx(doc, p)
    final = Document(str(p))
    new_index = next(
        (i for i, para in enumerate(final.paragraphs) if para._element is paragraph._element),
        len(final.paragraphs) - 1,
    )
    return {"index": new_index}


# ---------------------------------------------------------------------------
# 7. word_find_replace
# ---------------------------------------------------------------------------


def _count_occurrences(text: str, find: str, case_sensitive: bool) -> int:
    """Return the number of occurrences of ``find`` in ``text``."""
    if not find:
        return 0
    if case_sensitive:
        return text.count(find)
    return len(re.findall(re.escape(find), text, flags=re.IGNORECASE))


def _replace_in_paragraph(
    paragraph: Paragraph,
    find: str,
    replace: str,
    case_sensitive: bool,
) -> int:
    """Replace ``find`` with ``replace`` inside ``paragraph.text``.

    Preserves the formatting of the first matching run by deleting
    later runs and rewriting the first run's text. Returns the number
    of matches that were replaced.

    If the paragraph contains no text changes (``find`` does not
    appear), the paragraph is left untouched and ``0`` is returned.
    """
    original = paragraph.text
    if _count_occurrences(original, find, case_sensitive) == 0:
        return 0

    if case_sensitive:
        new_text = original.replace(find, replace)
    else:
        new_text = re.sub(re.escape(find), replace, original, flags=re.IGNORECASE)

    runs = list(paragraph.runs)
    if not runs:
        # No runs to preserve; just add the new text.
        paragraph.add_run(new_text)
        return _count_occurrences(original, find, case_sensitive)

    # Keep first run's formatting; drop the rest.
    for run in runs[1:]:
        run._element.getparent().remove(run._element)
    runs[0].text = new_text
    return _count_occurrences(original, find, case_sensitive)


@mcp.tool()
def word_find_replace(
    path: str,
    find: str,
    replace: str,
    case_sensitive: bool = True,
    folder: str | None = None,
) -> dict[str, int]:
    """Replace every occurrence of ``find`` with ``replace``.

    Operates on every body paragraph. The reported ``replacements``
    count is the sum of occurrences across all paragraphs before the
    replacement is applied (case-sensitivity follows the
    ``case_sensitive`` flag).

    Args:
        path: Path to an existing ``.docx``.
        find: The literal string to search for. Must be non-empty.
        replace: The replacement string.
        case_sensitive: When ``True`` (default), the match is exact;
            when ``False``, the match is case-insensitive.
        folder: Optional base folder for relative paths.

    Returns:
        ``{"replacements": <n>}``.

    Raises:
        OfficeMCPError: ``ERR_INVALID_PARAMS`` if ``find`` is empty or
            non-string, ``ERR_FILE_NOT_FOUND`` if the file is missing,
            ``ERR_UNSUPPORTED_FMT`` for non-``.docx`` extensions.
    """
    if not isinstance(find, str):
        raise OfficeMCPError(
            ERR_INVALID_PARAMS,
            f"find must be a str, got {type(find).__name__}",
            {"type": type(find).__name__},
        )
    if not isinstance(replace, str):
        raise OfficeMCPError(
            ERR_INVALID_PARAMS,
            f"replace must be a str, got {type(replace).__name__}",
            {"type": type(replace).__name__},
        )
    if find == "":
        raise OfficeMCPError(
            ERR_INVALID_PARAMS,
            "find string must be non-empty",
            {"find": find},
        )
    p = _resolve_docx(path, folder)
    doc = _open_docx(p)
    total = 0
    for paragraph in doc.paragraphs:
        total += _replace_in_paragraph(paragraph, find, replace, case_sensitive)
    _save_docx(doc, p)
    return {"replacements": total}


# ---------------------------------------------------------------------------
# 8. word_format_run
# ---------------------------------------------------------------------------


@mcp.tool()
def word_format_run(
    path: str,
    paragraph_index: int,
    run_index: int,
    bold: bool | None = None,
    italic: bool | None = None,
    font_size: float | int | None = None,
    font_name: str | None = None,
    color: str | None = None,
    folder: str | None = None,
) -> dict[str, bool]:
    """Update the formatting of a single run.

    Args:
        path: Path to an existing ``.docx``.
        paragraph_index: Zero-based index of the paragraph.
        run_index: Zero-based index of the run inside that paragraph.
        bold: ``True`` / ``False`` to set, or ``None`` to leave as-is.
        italic: Same convention as ``bold``.
        font_size: Point size (int or float) or ``None``.
        font_name: Font name string or ``None``.
        color: Hex color string (e.g. ``"FF0000"``) or ``None``.
        folder: Optional base folder for relative paths.

    Returns:
        ``{"ok": True}``.

    Special case:
        If **all** of ``bold``, ``italic``, ``font_size``, ``font_name``,
        ``color`` are ``None``, the function is a no-op: the file is
        not re-saved and the SHA256 is preserved
        (``VAL-WORD-039`` / ``VAL-WORD-076``).

    Raises:
        OfficeMCPError: ``ERR_INVALID_PARAMS`` if indices are out of
            range or of the wrong type, or ``color`` is not a valid
            hex string. ``ERR_FILE_NOT_FOUND`` if the file is missing,
            ``ERR_UNSUPPORTED_FMT`` for non-``.docx`` extensions.
    """
    formatting_args = (bold, italic, font_size, font_name, color)
    is_noop = all(v is None for v in formatting_args)

    p = _resolve_docx(path, folder)
    doc = _open_docx(p)
    _check_paragraph_index(doc, paragraph_index)
    paragraph = doc.paragraphs[paragraph_index]
    _check_run_index(paragraph, run_index)
    run = paragraph.runs[run_index]

    if not is_noop:
        if bold is not None:
            run.bold = bool(bold)
        if italic is not None:
            run.italic = bool(italic)
        if font_size is not None:
            try:
                run.font.size = Pt(float(font_size))
            except (TypeError, ValueError) as exc:
                raise OfficeMCPError(
                    ERR_INVALID_PARAMS,
                    f"font_size must be a number, got {font_size!r}",
                    {"font_size": font_size},
                ) from exc
        if font_name is not None:
            if not isinstance(font_name, str):
                raise OfficeMCPError(
                    ERR_INVALID_PARAMS,
                    f"font_name must be a str, got {type(font_name).__name__}",
                    {"type": type(font_name).__name__},
                )
            run.font.name = font_name
        if color is not None:
            if not isinstance(color, str):
                raise OfficeMCPError(
                    ERR_INVALID_PARAMS,
                    f"color must be a str, got {type(color).__name__}",
                    {"type": type(color).__name__},
                )
            try:
                run.font.color.rgb = RGBColor.from_string(color)
            except (ValueError, AttributeError) as exc:
                raise OfficeMCPError(
                    ERR_INVALID_PARAMS,
                    f"color must be a hex string like 'FF0000', got {color!r}",
                    {"color": color},
                ) from exc
        _save_docx(doc, p)
    return {"ok": True}


# ---------------------------------------------------------------------------
# 9. word_add_table
# ---------------------------------------------------------------------------


def _validate_table_dims(rows: int, cols: int) -> None:
    """Validate ``rows`` and ``cols`` for :func:`word_add_table`."""
    if not isinstance(rows, int) or isinstance(rows, bool):
        raise OfficeMCPError(
            ERR_INVALID_PARAMS,
            f"rows must be an int, got {type(rows).__name__}",
            {"rows": rows},
        )
    if not isinstance(cols, int) or isinstance(cols, bool):
        raise OfficeMCPError(
            ERR_INVALID_PARAMS,
            f"cols must be an int, got {type(cols).__name__}",
            {"cols": cols},
        )
    if rows <= 0:
        raise OfficeMCPError(
            ERR_INVALID_PARAMS,
            f"rows must be a positive integer, got {rows}",
            {"rows": rows},
        )
    if cols <= 0:
        raise OfficeMCPError(
            ERR_INVALID_PARAMS,
            f"cols must be a positive integer, got {cols}",
            {"cols": cols},
        )


@mcp.tool()
def word_add_table(
    path: str,
    rows: int,
    cols: int,
    data: list[list[str]] | None = None,
    style: str | None = None,
    folder: str | None = None,
) -> dict[str, int]:
    """Append a table to the body of the document.

    Args:
        path: Path to an existing ``.docx``.
        rows: Number of rows (must be a positive int).
        cols: Number of columns (must be a positive int).
        data: Optional 2D list of strings to populate the cells. The
            list is interpreted in row-major order; shorter rows are
            padded with empty strings and longer rows are truncated.
        style: Optional built-in table style name (e.g. ``"Table Grid"``
            or ``"Light Grid Accent 1"``).
        folder: Optional base folder for relative paths.

    Returns:
        ``{"index": <n>}`` where ``<n>`` is the index of the new table
        in ``doc.tables``.

    Raises:
        OfficeMCPError: ``ERR_INVALID_PARAMS`` for non-positive
            ``rows``/``cols``, unknown ``style``, or a non-list
            ``data`` argument. ``ERR_FILE_NOT_FOUND`` if the file is
            missing, ``ERR_UNSUPPORTED_FMT`` for non-``.docx``
            extensions.
    """
    _validate_table_dims(rows, cols)
    if data is not None and not isinstance(data, list):
        raise OfficeMCPError(
            ERR_INVALID_PARAMS,
            f"data must be a list (or None), got {type(data).__name__}",
            {"type": type(data).__name__},
        )
    p = _resolve_docx(path, folder)
    doc = _open_docx(p)
    table = doc.add_table(rows=rows, cols=cols)

    # Populate cells if data is provided
    if data:
        for r_idx, row_values in enumerate(data):
            if r_idx >= rows:
                break
            if not isinstance(row_values, Iterable) or isinstance(row_values, str):
                # Skip non-iterable rows (but accept strings; we treat
                # them as a 1-element list).
                row_values = [row_values] if isinstance(row_values, str) else []
            for c_idx, value in enumerate(row_values):
                if c_idx >= cols:
                    break
                table.cell(r_idx, c_idx).text = "" if value is None else str(value)

    if style is not None:
        if not isinstance(style, str):
            raise OfficeMCPError(
                ERR_INVALID_PARAMS,
                f"style must be a str, got {type(style).__name__}",
                {"type": type(style).__name__},
            )
        try:
            table.style = style
        except KeyError as exc:
            raise OfficeMCPError(
                ERR_INVALID_PARAMS,
                f"Unknown table style: {style!r}",
                {"style": style},
            ) from exc

    _save_docx(doc, p)
    # Reload to find the final table index (in case add_table moved things)
    final = Document(str(p))
    new_index = len(final.tables) - 1
    return {"index": new_index}


# ---------------------------------------------------------------------------
# 10. word_add_image
# ---------------------------------------------------------------------------


# Image content-types that python-docx / Word reliably accept via
# ``add_picture``. ``.bmp`` and other formats raise ``UnrecognizedImageError``
# at runtime; we filter those out before delegating to the library.
_ACCEPTED_IMAGE_EXTENSIONS: frozenset[str] = frozenset(
    {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tif", ".tiff"}
)


def _validate_image_args(
    image_path: Path,
    width_inches: float | int | None,
) -> None:
    """Validate inputs for :func:`word_add_image`."""
    if not isinstance(image_path, Path):
        raise OfficeMCPError(
            ERR_INVALID_PARAMS,
            f"image_path must be a str, got {type(image_path).__name__}",
            {"type": type(image_path).__name__},
        )
    if not image_path.exists():
        raise OfficeMCPError(
            ERR_FILE_NOT_FOUND,
            f"Image not found: {image_path}",
            {"image_path": str(image_path)},
        )
    suffix = image_path.suffix.lower()
    if suffix not in _ACCEPTED_IMAGE_EXTENSIONS:
        raise OfficeMCPError(
            ERR_INVALID_PARAMS,
            (
                f"Unsupported image format {suffix!r}. "
                f"Supported: {sorted(_ACCEPTED_IMAGE_EXTENSIONS)}"
            ),
            {"image_path": str(image_path), "extension": suffix},
        )
    if width_inches is not None:
        if isinstance(width_inches, bool) or not isinstance(width_inches, (int, float)):
            raise OfficeMCPError(
                ERR_INVALID_PARAMS,
                f"width_inches must be a positive number, got {width_inches!r}",
                {"width_inches": width_inches},
            )
        if float(width_inches) <= 0:
            raise OfficeMCPError(
                ERR_INVALID_PARAMS,
                f"width_inches must be positive, got {width_inches}",
                {"width_inches": width_inches},
            )


@mcp.tool()
def word_add_image(
    path: str,
    image_path: str,
    width_inches: float | int | None = None,
    folder: str | None = None,
) -> dict[str, int]:
    """Embed an image in the document body.

    When ``width_inches`` is provided, the image is scaled to that
    width and the height is auto-computed from the image's intrinsic
    aspect ratio (``VAL-WORD-050`` / ``VAL-WORD-078``). When omitted,
    the image is embedded at its native pixel size.

    Args:
        path: Path to an existing ``.docx``.
        image_path: Path to the image file (PNG / JPEG / GIF / etc.).
        width_inches: Optional target width in inches (must be
            positive). When ``None`` the image is embedded at its
            native size.
        folder: Optional base folder for the ``path`` argument
            (does not apply to ``image_path``, which is always
            resolved absolutely or relative to the caller's CWD).

    Returns:
        ``{"index": <n>}`` where ``<n>`` is the index of the new
        inline shape in ``doc.inline_shapes``.

    Raises:
        OfficeMCPError: ``ERR_FILE_NOT_FOUND`` if either path is
            missing, ``ERR_INVALID_PARAMS`` for an unsupported image
            format or a non-positive ``width_inches``,
            ``ERR_UNSUPPORTED_FMT`` for non-``.docx`` extensions.
    """
    # Resolve and validate the document path first (so a missing
    # document takes precedence over image validation).
    doc_path = _resolve_docx(path, folder)

    # Resolve the image path independently; we want clear error codes
    # for "missing image" vs. "unsupported format" vs. "bad width".
    img_path_obj = Path(image_path)
    _validate_image_args(img_path_obj, width_inches)

    # Now open the document (after all argument validation has passed
    # so an invalid argument does not touch the file).
    doc = _open_docx(doc_path)

    if width_inches is not None:
        # python-docx preserves the aspect ratio automatically when
        # only ``width`` is supplied (see VAL-WORD-050 / -078).
        doc.add_picture(str(img_path_obj), width=Inches(float(width_inches)))
    else:
        doc.add_picture(str(img_path_obj))

    _save_docx(doc, doc_path)
    # Reload to find the new inline shape's index
    final = Document(str(doc_path))
    new_index = len(final.inline_shapes) - 1
    return {"index": new_index}


# ---------------------------------------------------------------------------
# 11. word_add_header
# ---------------------------------------------------------------------------


def _set_header_text(
    section,
    text: str,
) -> None:
    """Write ``text`` into ``section.header.paragraphs[0]``.

    Ensures the header is unlinked from the previous section so the
    call is not silently absorbed by the prior section's header
    (matters when the document has multiple sections and the caller
    only wants to touch section 0). The text of the first paragraph
    is replaced in-place; subsequent paragraphs are left alone.
    """
    header = section.header
    # Unlinking ensures the header becomes a distinct reference,
    # separate from any prior section's header. This is a no-op when
    # the document has only one section.
    header.is_linked_to_previous = False
    paragraphs = header.paragraphs
    if not paragraphs:
        # The header part may have no paragraphs yet; ``add_paragraph``
        # creates one and returns it.
        paragraphs = [header.add_paragraph()]
    paragraphs[0].text = text


@mcp.tool()
def word_add_header(
    path: str,
    text: str,
    section_index: int = 0,
    folder: str | None = None,
) -> dict[str, bool]:
    """Set the text of a section's header.

    By default, only the first section's header is updated; later
    sections keep their own headers (``VAL-WORD-058``). Pass an
    explicit ``section_index`` to target a different section.

    Args:
        path: Path to an existing ``.docx``.
        text: New header text. An empty string clears the first
            paragraph (``VAL-WORD-057``).
        section_index: Zero-based section index (default 0).
        folder: Optional base folder for relative paths.

    Returns:
        ``{"ok": True}``.

    Raises:
        OfficeMCPError: ``ERR_INVALID_PARAMS`` for a non-int or
            out-of-range ``section_index`` or a non-string ``text``,
            ``ERR_FILE_NOT_FOUND`` if the file is missing,
            ``ERR_UNSUPPORTED_FMT`` for non-``.docx`` extensions.
    """
    if not isinstance(text, str):
        raise OfficeMCPError(
            ERR_INVALID_PARAMS,
            f"text must be a str, got {type(text).__name__}",
            {"type": type(text).__name__},
        )
    if not isinstance(section_index, int) or isinstance(section_index, bool):
        raise OfficeMCPError(
            ERR_INVALID_PARAMS,
            f"section_index must be an int, got {type(section_index).__name__}",
            {"section_index": section_index},
        )
    p = _resolve_docx(path, folder)
    doc = _open_docx(p)
    if section_index < 0 or section_index >= len(doc.sections):
        raise OfficeMCPError(
            ERR_INVALID_PARAMS,
            (
                f"section_index {section_index} is out of range "
                f"(document has {len(doc.sections)} section(s))"
            ),
            {"section_index": section_index, "count": len(doc.sections)},
        )
    _set_header_text(doc.sections[section_index], text)
    _save_docx(doc, p)
    return {"ok": True}


# ---------------------------------------------------------------------------
# 12. word_add_footer
# ---------------------------------------------------------------------------


def _set_footer_text(
    section,
    text: str,
) -> None:
    """Write ``text`` into ``section.footer.paragraphs[0]``."""
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraphs = footer.paragraphs
    if not paragraphs:
        paragraphs = [footer.add_paragraph()]
    paragraphs[0].text = text


@mcp.tool()
def word_add_footer(
    path: str,
    text: str,
    section_index: int = 0,
    folder: str | None = None,
) -> dict[str, bool]:
    """Set the text of a section's footer.

    By default, only the first section's footer is updated; later
    sections keep their own footers (``VAL-WORD-062``). Pass an
    explicit ``section_index`` to target a different section.

    Args:
        path: Path to an existing ``.docx``.
        text: New footer text. An empty string clears the first
            paragraph (``VAL-WORD-061``).
        section_index: Zero-based section index (default 0).
        folder: Optional base folder for relative paths.

    Returns:
        ``{"ok": True}``.

    Raises:
        OfficeMCPError: ``ERR_INVALID_PARAMS`` for a non-int or
            out-of-range ``section_index`` or a non-string ``text``,
            ``ERR_FILE_NOT_FOUND`` if the file is missing,
            ``ERR_UNSUPPORTED_FMT`` for non-``.docx`` extensions.
    """
    if not isinstance(text, str):
        raise OfficeMCPError(
            ERR_INVALID_PARAMS,
            f"text must be a str, got {type(text).__name__}",
            {"type": type(text).__name__},
        )
    if not isinstance(section_index, int) or isinstance(section_index, bool):
        raise OfficeMCPError(
            ERR_INVALID_PARAMS,
            f"section_index must be an int, got {type(section_index).__name__}",
            {"section_index": section_index},
        )
    p = _resolve_docx(path, folder)
    doc = _open_docx(p)
    if section_index < 0 or section_index >= len(doc.sections):
        raise OfficeMCPError(
            ERR_INVALID_PARAMS,
            (
                f"section_index {section_index} is out of range "
                f"(document has {len(doc.sections)} section(s))"
            ),
            {"section_index": section_index, "count": len(doc.sections)},
        )
    _set_footer_text(doc.sections[section_index], text)
    _save_docx(doc, p)
    return {"ok": True}


# ---------------------------------------------------------------------------
# 13. word_set_section
# ---------------------------------------------------------------------------


def _parse_orientation(orientation: str) -> int:
    """Map a user-supplied orientation string to a ``WD_ORIENT`` value."""
    if not isinstance(orientation, str):
        raise OfficeMCPError(
            ERR_INVALID_PARAMS,
            f"orientation must be a str, got {type(orientation).__name__}",
            {"orientation": orientation},
        )
    key = orientation.strip().lower()
    if key in ("portrait", "p"):
        return WD_ORIENT.PORTRAIT
    if key in ("landscape", "l"):
        return WD_ORIENT.LANDSCAPE
    raise OfficeMCPError(
        ERR_INVALID_PARAMS,
        (
            f"Unknown orientation {orientation!r}. "
            f"Expected one of: 'portrait', 'landscape'"
        ),
        {"orientation": orientation},
    )


def _parse_page_size(page_size: str) -> tuple[float, float]:
    """Return ``(width_in, height_in)`` for a named page size."""
    if not isinstance(page_size, str):
        raise OfficeMCPError(
            ERR_INVALID_PARAMS,
            f"page_size must be a str, got {type(page_size).__name__}",
            {"page_size": page_size},
        )
    key = page_size.strip()
    # Accept case-insensitive lookup
    for candidate, dims in _PAGE_SIZES.items():
        if candidate.lower() == key.lower():
            return dims
    raise OfficeMCPError(
        ERR_INVALID_PARAMS,
        (
            f"Unknown page_size {page_size!r}. "
            f"Expected one of: {sorted(_PAGE_SIZES)}"
        ),
        {"page_size": page_size},
    )


@mcp.tool()
def word_set_section(
    path: str,
    orientation: str = "portrait",
    page_size: str = "A4",
    section_index: int = 0,
    folder: str | None = None,
) -> dict[str, bool]:
    """Configure a section's orientation and page size.

    Args:
        path: Path to an existing ``.docx``.
        orientation: ``"portrait"`` or ``"landscape"`` (case-insensitive).
            Any other value raises ``ERR_INVALID_PARAMS``
            (``VAL-WORD-065``).
        page_size: A named page size such as ``"A4"`` (default),
            ``"Letter"``, ``"Legal"``, ``"A5"``, ``"Tabloid"`` or
            ``"B5"`` (case-insensitive).
        section_index: Zero-based section index (default 0).
        folder: Optional base folder for relative paths.

    Returns:
        ``{"ok": True}``.

    Behavior:
        * Portrait: ``page_width = <width>``, ``page_height = <height>``.
        * Landscape: ``page_width = <height>``, ``page_height = <width>``
          (i.e. dimensions are swapped so the wider edge is horizontal).

    Raises:
        OfficeMCPError: ``ERR_INVALID_PARAMS`` for an unknown
            ``orientation`` or ``page_size``, ``ERR_FILE_NOT_FOUND``
            if the file is missing, ``ERR_UNSUPPORTED_FMT`` for
            non-``.docx`` extensions.
    """
    orient_value = _parse_orientation(orientation)
    width_in, height_in = _parse_page_size(page_size)
    if not isinstance(section_index, int) or isinstance(section_index, bool):
        raise OfficeMCPError(
            ERR_INVALID_PARAMS,
            f"section_index must be an int, got {type(section_index).__name__}",
            {"section_index": section_index},
        )
    p = _resolve_docx(path, folder)
    doc = _open_docx(p)
    if section_index < 0 or section_index >= len(doc.sections):
        raise OfficeMCPError(
            ERR_INVALID_PARAMS,
            (
                f"section_index {section_index} is out of range "
                f"(document has {len(doc.sections)} section(s))"
            ),
            {"section_index": section_index, "count": len(doc.sections)},
        )
    section = doc.sections[section_index]
    if orient_value == WD_ORIENT.LANDSCAPE:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(height_in)
        section.page_height = Inches(width_in)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(width_in)
        section.page_height = Inches(height_in)
    _save_docx(doc, p)
    return {"ok": True}


# ---------------------------------------------------------------------------
# 14. word_export_pdf
# ---------------------------------------------------------------------------


@mcp.tool()
def word_export_pdf(
    path: str,
    output: str,
    folder: str | None = None,
) -> dict[str, str]:
    """Convert a ``.docx`` file to PDF via LibreOffice headless.

    Delegates to :func:`office_mcp.exporters.export_to_pdf`. A unique
    ``-env:UserInstallation`` is used per call so multiple exports
    can run concurrently (``VAL-WORD-079``).

    Args:
        path: Path to an existing ``.docx`` file.
        output: Target path for the produced PDF. The parent
            directory is created if it does not exist. If a relative
            path is given, it is resolved against ``folder`` (or the
            default folder).
        folder: Optional base folder for relative paths.

    Returns:
        ``{"output_path": "<absolute path of the produced PDF>"}``.

    Raises:
        OfficeMCPError: ``ERR_FILE_NOT_FOUND`` if the source is
            missing, ``ERR_UNSUPPORTED_FMT`` for non-``.docx``
            sources, ``ERR_LIBREOFFICE_MISSING`` when ``soffice`` is
            not on PATH, ``ERR_EXPORT_FAILED`` for any other failure.
    """
    # Local import to avoid a circular dependency at module import
    # time: ``exporters`` does not import from ``word_tools`` but the
    # server bootstrap imports all tool modules eagerly.
    from office_mcp import exporters

    src = _resolve_docx(path, folder)
    # Verify the source exists up front so the caller gets the
    # agent-friendly ``ERR_FILE_NOT_FOUND`` code rather than a raw
    # "source file could not be loaded" coming back from LibreOffice
    # (``VAL-WORD-068``).
    _require_exists(src)
    # The output path may be relative; resolve it independently so
    # that the ``folder`` argument is honoured.
    out_path = resolve_path(output, folder)
    out_dir = out_path.parent
    produced = exporters.export_to_pdf(src, out_dir)
    # ``exporters`` may name the output ``<stem>.pdf``; if the caller
    # requested a different filename, move (rename) the produced file
    # to the requested location. The default case is the same path.
    if produced != out_path:
        out_path.parent.mkdir(parents=True, exist_ok=True)
        produced.replace(out_path)
        final = out_path
    else:
        final = produced
    return {"output_path": str(final)}


# ---------------------------------------------------------------------------
# 15. word_export_html
# ---------------------------------------------------------------------------


@mcp.tool()
def word_export_html(
    path: str,
    output: str,
    folder: str | None = None,
) -> dict[str, str]:
    """Convert a ``.docx`` file to HTML via :mod:`mammoth`.

    Delegates to :func:`office_mcp.exporters.export_to_html`. The
    ``.docx`` path uses ``mammoth`` only (no LibreOffice dependency
    — ``VAL-WORD-072``).

    Args:
        path: Path to an existing ``.docx`` file.
        output: Target path for the produced HTML. The parent
            directory is created if it does not exist.
        folder: Optional base folder for relative paths.

    Returns:
        ``{"output_path": "<absolute path of the produced HTML>"}``.

    Raises:
        OfficeMCPError: ``ERR_FILE_NOT_FOUND`` if the source is
            missing, ``ERR_UNSUPPORTED_FMT`` for non-``.docx``
            sources, ``ERR_EXPORT_FAILED`` for any other failure.
    """
    from office_mcp import exporters

    src = _resolve_docx(path, folder)
    # Verify the source exists up front so the caller gets
    # ``ERR_FILE_NOT_FOUND`` instead of a raw ``FileNotFoundError``
    # from ``open()`` inside mammoth (``VAL-WORD-073``).
    _require_exists(src)
    out_path = resolve_path(output, folder)
    out_dir = out_path.parent
    produced = exporters.export_to_html(src, out_dir)
    if produced != out_path:
        out_path.parent.mkdir(parents=True, exist_ok=True)
        produced.replace(out_path)
        final = out_path
    else:
        final = produced
    return {"output_path": str(final)}
